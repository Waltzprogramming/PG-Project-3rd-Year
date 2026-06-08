from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


ROOT = Path(r"C:\Proyectos\PROYECTO")
OUT = ROOT / "artifacts" / "mapa4_desglose_tecnico.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1E1E1E")
    p = cell.paragraphs[0]
    for idx, line in enumerate(code.strip("\n").splitlines()):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(240, 240, 240)
        if idx < len(code.strip("\n").splitlines()) - 1:
            p.add_run("\n")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = "Normal"
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(11)
    return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Normal"].font.size = Pt(11)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Heading 1"].font.size = Pt(18)
styles["Heading 1"].font.color.rgb = RGBColor(18, 64, 120)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Heading 2"].font.size = Pt(14)
styles["Heading 2"].font.color.rgb = RGBColor(30, 30, 30)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("desglose técnico de mapa 4")
r.bold = True
r.font.name = "Arial"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(20, 57, 102)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("explicación de luz, soles, escudo y enemigos usando fragmentos reales del proyecto")
r.font.name = "Arial"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(90, 90, 90)

sections = [
    {
        "title": "1. sistema de luz",
        "line": "líneas iniciales: Map4.cpp 1115 y main.cpp 2106",
        "body": (
            "La funcionalidad de la luz se coordina desde la función renderMapa4(...), que es la función principal del frame de Mapa 4. "
            "Dentro de esa función, flashlightAnchor es una variable local de tipo glm::vec3 y mapa4.lightEnergy es un campo del struct Mapa4Runtime. "
            "La función renderMapa4(...) calcula la posición actual del jugador, toma la energía disponible y luego llama a la función uploadCommonSceneUniforms(...), "
            "que es la encargada de enviar esa información al shader. En términos técnicos, la luz del jugador se comporta como una point light dinámica: su posición se actualiza cada frame, "
            "su intensidad depende del ratio de energía y su radio de influencia se reduce o aumenta según playerLightRatio."
        ),
        "code": """const glm::vec3 flashlightAnchor = mapa4.player.position();
uploadCommonSceneUniforms(sceneShader, mapa4.environment, gameplayCameraPosition, view, projection, now, &flashlightAnchor, mapa4.lightEnergy / Map4LightEnergyMaximum, nullptr);""",
        "body2": (
            "La otra parte importante está en uploadCommonSceneUniforms(...), que también es una función. Ahí, playerLightPosition y playerLightRatio son parámetros, "
            "prefix es una variable local para construir el nombre del uniform, e intensity y radius son variables locales que definen la potencia real de la luz. "
            "Eso significa que cuando hables de esa parte del código, lo correcto es decir que renderMapa4(...) llama a otra función que crea una luz extra del shader usando una variable local "
            "como ancla y un campo del runtime como fuente de energía."
        ),
        "code2": """if (playerLightPosition && nextLightIndex < count + extraCount) {
    const std::string prefix = "uPointLights[" + std::to_string(nextLightIndex) + "]";
    shader.setVec3(prefix + ".position", *playerLightPosition + glm::vec3(0.0f, 0.46f, currentMode == PlayMode::Mode2D ? 0.36f : 0.18f));
    shader.setVec3(prefix + ".color", glm::vec3(0.94f, 0.88f, 0.72f));
    const float lightScale = std::clamp(playerLightRatio, 0.0f, 1.0f);
    const float intensity = lightScale * 1.18f;
    const float radius = lightScale * 3.45f;
    shader.setFloat(prefix + ".intensity", intensity);
    shader.setFloat(prefix + ".radius", radius);
}""",
    },
    {
        "title": "2. sistema de soles",
        "line": "líneas iniciales: Map4.cpp 721, 737 y 834",
        "body": (
            "Los soles están administrados por la clase Map4LightManager. En este caso, initialize(), reset(...) y update(...) son métodos de esa clase. "
            "La inicialización carga el recurso visual del Sol, reset(...) vuelve a poblar sus posiciones en superficies alcanzables, y update(...) resuelve tanto el consumo de energía como la recolección y el respawn. "
            "Dentro de reset(...), m_pickups es un campo privado de la clase y SurfaceCandidate es una estructura local auxiliar usada para filtrar suelos válidos del escenario. "
            "Técnicamente, el sistema recorre environment.collisionPreview(), detecta plataformas con suficiente área útil y deja los Soles solamente en zonas jugables cercanas a la altura del jugador."
        ),
        "code": """void Map4LightManager::reset(const Environment& environment, const glm::vec3& playerSpawn) {
    m_pickups.clear();
    struct SurfaceCandidate {
        Bounds bounds;
        float top{0.0f};
        float area{0.0f};
    };

    std::vector<SurfaceCandidate> surfaces;
    for (const Bounds& collider : environment.collisionPreview()) {
        const float top = collider.center.y + collider.halfExtent.y;
        const float area = (collider.halfExtent.x * 2.0f) * (collider.halfExtent.z * 2.0f);
        const bool floorLike = collider.halfExtent.y <= 0.32f && area >= 0.85f && collider.halfExtent.x >= 0.38f && collider.halfExtent.z >= 0.38f;
        const bool reachableHeight = top >= playerSpawn.y - 0.25f && top <= playerSpawn.y + 1.25f;
        if (floorLike && reachableHeight) {
            surfaces.push_back({collider, top, area});
        }
    }""",
        "body2": (
            "La parte de gameplay importante está en update(...). Ahí, energySeconds es un parámetro por referencia, pickup es una variable de iteración del for, "
            "y pickup.available y pickup.respawnAt son campos del struct SunPickup. Cuando el jugador toca un Sol, el pickup se desactiva, se agenda su respawn y se recarga la energía. "
            "La forma técnica correcta de decirlo es que el método update(...) modifica un parámetro por referencia que representa la barra de luz, usando el estado interno de cada pickup."
        ),
        "code2": """void Map4LightManager::update(Player& player, float timeSeconds, float deltaTime, float& energySeconds) {
    energySeconds = std::max(0.0f, energySeconds - deltaTime);
    const Bounds playerBounds = player.bounds();
    for (SunPickup& pickup : m_pickups) {
        if (!pickup.available) {
            if (timeSeconds >= pickup.respawnAt) {
                pickup.available = true;
            }
            continue;
        }

        if (map4BoundsIntersect(playerBounds, sunBounds(pickup))) {
            pickup.available = false;
            pickup.respawnAt = timeSeconds + Map4LightRespawnTime;
            energySeconds = std::min(Map4LightEnergyMaximum, energySeconds + Map4LightEnergyMaximum);
        }
    }
}""",
    },
    {
        "title": "3. sistema de escudo",
        "line": "líneas iniciales: Map4.cpp 80, 217 y 1115",
        "body": (
            "El escudo tiene una parte lógica y una parte visual. La parte lógica vive dentro de la función renderMapa4(...), donde shieldDown y shieldPressed son variables locales usadas para leer la entrada, "
            "mientras que mapa4.shieldActive y mapa4.shieldTimer son campos del struct Mapa4Runtime que conservan el estado del escudo entre frames. "
            "Cada vez que se presiona E, el sistema cambia el estado del escudo y le asigna una duración fija de un segundo. Luego, en cada iteración del frame, ese temporizador se reduce con frameDelta hasta apagarse."
        ),
        "code": """const bool shieldDown = glfwGetKey(window, GLFW_KEY_E) == GLFW_PRESS;
const bool shieldPressed = shieldDown && !lastShieldKey;
lastShieldKey = shieldDown;
if (shieldPressed) {
    if (mapa4.shieldActive) {
        mapa4.shieldActive = false;
        mapa4.shieldTimer = 0.0f;
    } else {
        mapa4.shieldActive = true;
        mapa4.shieldTimer = 1.0f;
    }
}
if (mapa4.shieldActive) {
    mapa4.shieldTimer = std::max(0.0f, mapa4.shieldTimer - frameDelta);
    if (mapa4.shieldTimer <= 0.0f) {
        mapa4.shieldActive = false;
    }
}""",
        "body2": (
            "La parte visual se construye en la función renderMapa4Shield(...), mientras que createMapa4ShieldRingMesh() es otra función auxiliar que genera una malla procedural. "
            "Dentro de renderMapa4Shield(...), shieldRing es una variable estática local, ringMaterial es una variable local de tipo Material y frontRing, sideRing y horizontalRing son matrices locales de transformación. "
            "Eso permite dibujar varios anillos cruzados alrededor del jugador sin cargar un modelo externo. Además, la defensa real se resuelve en updateMapa4Projectiles(...): si el proyectil enemigo toca al jugador y el campo mapa4.shieldActive está activo, el daño no se aplica."
        ),
        "code2": """void renderMapa4Shield(const Shader& shader, const Player& player, float timeSeconds) {
    static Mesh shieldRing = createMapa4ShieldRingMesh();

    const float pulse = 0.96f + 0.04f * std::sin(timeSeconds * 9.0f);
    Material ringMaterial;
    ringMaterial.baseColor = {0.22f, 1.00f, 0.94f};
    ringMaterial.emissive = {0.02f, 0.08f, 0.08f};
    ringMaterial.roughness = 0.95f;
    ringMaterial.fogAmount = 0.02f;
    ringMaterial.opacity = 0.16f;

    const glm::vec3 center = player.position() + glm::vec3(0.0f, 0.58f, 0.0f);
    glm::mat4 frontRing(1.0f);
    frontRing = glm::translate(frontRing, center);
    frontRing = glm::scale(frontRing, glm::vec3(0.88f, 0.94f, 0.88f) * pulse);""",
    },
    {
        "title": "4. sistema de enemigos",
        "line": "líneas iniciales: Map4.cpp 334, 346, 380 y 462",
        "body": (
            "Los enemigos se controlan desde la clase SimpleEnemyManager. En esa clase, initialize(), reset(...), update(...) y render(...) son métodos. "
            "m_models y m_enemies son campos privados del manager, mientras que Enemy es una estructura interna que guarda el estado individual de cada enemigo, incluyendo posición, patrulla, cooldown, rangos de detección y vida. "
            "En reset(...), anchors es una constante local y enemy es una variable local que se va llenando antes de insertarse en el vector m_enemies. "
            "El objetivo de ese método es repartir enemigos sobre pisos válidos, alejados del spawn del jugador y con parámetros de patrulla ligeramente distintos."
        ),
        "code": """void SimpleEnemyManager::reset(const Environment& environment, const glm::vec3& playerSpawn) {
    m_enemies.clear();
    const glm::vec3 worldMin = environment.worldMin();
    const glm::vec3 worldMax = environment.worldMax();
    const glm::vec2 center((worldMin.x + worldMax.x) * 0.5f, (worldMin.z + worldMax.z) * 0.5f);
    const std::array<glm::vec2, 6> anchors = {
        center + glm::vec2(-7.0f, -5.4f),
        center + glm::vec2(7.0f, -4.8f),
        center + glm::vec2(-6.1f, 5.8f),
        center + glm::vec2(6.3f, 5.6f),
        center + glm::vec2(-1.8f, -6.2f),
        center + glm::vec2(2.4f, 6.4f)
    };

    for (size_t index = 0; index < anchors.size(); ++index) {
        Enemy enemy;
        enemy.modelIndex = static_cast<int>(index % m_models.size());
        enemy.position = findSpawnPosition(environment, playerSpawn, anchors[index]);""",
        "body2": (
            "La IA principal está en update(...). Ahí, player, environment, deltaTime, timeSeconds y projectiles son parámetros. "
            "Dentro del método, dt, toPlayer y distance son variables locales que se recalculan por enemigo. "
            "Si el jugador está lejos, el enemigo se salta para ahorrar costo; si entra a rango, patrulla y luego gira hacia el jugador. Si también entra al rango de ataque y su cooldown ya venció, se genera un nuevo proyectil enemigo en el vector projectiles. "
            "Técnicamente, no hay navegación compleja ni pathfinding: el sistema usa patrulla oscilatoria, chequeo de distancia, colisión simple y spawn de proyectiles."
        ),
        "code2": """bool SimpleEnemyManager::update(const Player& player, const Environment& environment, float deltaTime, float timeSeconds, std::vector<Mapa4Projectile>& projectiles) {
    bool damagedPlayer = false;
    const Bounds playerBounds = player.bounds();
    const glm::vec3 playerPosition = player.position();
    const float dt = std::clamp(deltaTime, 0.0f, 1.0f / 30.0f);

    for (Enemy& enemy : m_enemies) {
        if (!enemy.alive) {
            continue;
        }
        enemy.attackCooldown = std::max(0.0f, enemy.attackCooldown - dt);
        glm::vec3 toPlayer = playerPosition - enemy.position;
        toPlayer.y = 0.0f;
        const float distance = glm::length(toPlayer);

        if (distance <= enemy.detectionRange && distance > 0.001f) {
            enemy.yaw = std::atan2(toPlayer.x, toPlayer.z);
            if (distance <= enemy.attackRange && enemy.attackCooldown <= 0.0f) {
                const glm::vec3 shotDirection = glm::normalize(glm::vec3(toPlayer.x, 0.0f, toPlayer.z));
                projectiles.push_back({
                    enemy.position + glm::vec3(shotDirection.x * 0.45f, 0.62f, shotDirection.z * 0.45f),
                    shotDirection * Map4EnemyProjectileSpeed,
                    Map4ProjectileLifetime,
                    1,
                    true
                });""",
    },
]

for section in sections:
    add_heading(doc, section["title"], 1)
    add_body(doc, section["line"])
    add_body(doc, section["body"])
    add_code_block(doc, section["code"])
    add_body(doc, section["body2"])
    add_code_block(doc, section["code2"])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
